#!/usr/bin/env python3
"""Regression checks for structured Skills replacement."""
from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from resume_patcher import insert_skill_paragraph_after, set_skills_text_keep_label_format


MASTER_NAME = "Caleb Miller Master Resume.docx"


def run_patcher(src: Path, replacements: Path, out: Path) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [
            sys.executable,
            str(REPO_ROOT / "resume_patcher.py"),
            "--src",
            str(src),
            "--replacements",
            str(replacements),
            "--out",
            str(out),
        ],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
    )


def write_replacements(path: Path, skills) -> None:
    path.write_text(json.dumps({"skills": skills}, indent=2), encoding="utf-8")


def find_section_bounds(doc: Document, heading: str = "SKILLS") -> tuple[int, int]:
    start = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip().upper() == heading)
    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        if paragraph.text.strip() and paragraph.style is not None and paragraph.style.name.lower() == "heading 1":
            end = i
            break
    return start, end


def skill_paragraphs(doc: Document) -> list:
    start, end = find_section_bounds(doc)
    return [paragraph for paragraph in doc.paragraphs[start + 1 : end] if paragraph.text.strip()]


def run_format(run) -> tuple[bool | None, bool | None]:
    return run.bold, run.italic


def first_run_with(paragraph, needle: str):
    for run in paragraph.runs:
        if needle in run.text:
            return run
    raise AssertionError(f"Could not find run containing {needle!r} in {paragraph.text!r}")


def make_multi_paragraph_skills_source(path: Path) -> None:
    doc = Document(str(REPO_ROOT / MASTER_NAME))
    source = skill_paragraphs(doc)[0]
    template = Paragraph(deepcopy(source._p), source._parent)
    lines = [line for line in source.text.splitlines() if line.strip()]

    set_skills_text_keep_label_format(source, lines[0], template)
    cursor = source
    for line in lines[1:2]:
        cursor = insert_skill_paragraph_after(cursor, line, template)

    doc.save(str(path))


def assert_single_paragraph_structured_skills(tmp_path: Path) -> None:
    src = REPO_ROOT / MASTER_NAME
    source_doc = Document(str(src))
    source_skill = skill_paragraphs(source_doc)[0]
    expected_style = source_skill.style.name
    expected_label_format = run_format(first_run_with(source_skill, "Support & Platform:"))
    expected_rest_format = run_format(first_run_with(source_skill, "technical support"))

    replacements = tmp_path / "single_replacements.json"
    out = tmp_path / "single_out.docx"
    write_replacements(
        replacements,
        [
            {
                "category": "Support & Customer Success",
                "items": ["enterprise SaaS support", "technical advising", "customer onboarding"],
            },
            {
                "category": "Systems & Infrastructure",
                "items": ["Windows Server", "Linux", "Active Directory", "DNS"],
            },
        ],
    )

    result = run_patcher(src, replacements, out)
    if result.returncode != 0:
        raise AssertionError(f"structured skills replacement failed:\n{result.stdout}\n{result.stderr}")

    output_doc = Document(str(out))
    output_skills = skill_paragraphs(output_doc)
    if len(output_skills) != 1:
        raise AssertionError(f"expected one Skills paragraph copied from master, got {len(output_skills)}")

    output_skill = output_skills[0]
    if output_skill.style.name != expected_style:
        raise AssertionError(f"Skills paragraph style changed: {output_skill.style.name!r} != {expected_style!r}")
    if "Support & Customer Success: enterprise SaaS support, technical advising, customer onboarding" not in output_skill.text:
        raise AssertionError("first structured skill category was not rendered")
    if "Systems & Infrastructure: Windows Server, Linux, Active Directory, DNS" not in output_skill.text:
        raise AssertionError("second structured skill category was not rendered")
    if "Observability & DevOps:" in output_skill.text:
        raise AssertionError("old extra Skills category was not removed from single-paragraph source")
    if run_format(first_run_with(output_skill, "Support & Customer Success:")) != expected_label_format:
        raise AssertionError("structured Skills label did not copy source label run formatting")
    if run_format(first_run_with(output_skill, "enterprise SaaS support")) != expected_rest_format:
        raise AssertionError("structured Skills item text did not copy source item run formatting")


def assert_multi_paragraph_add_and_remove(tmp_path: Path) -> None:
    src = tmp_path / MASTER_NAME
    make_multi_paragraph_skills_source(src)

    replacements_more = tmp_path / "multi_more_replacements.json"
    out_more = tmp_path / "multi_more_out.docx"
    write_replacements(
        replacements_more,
        [
            {"category": "One", "items": ["alpha"]},
            {"category": "Two", "items": ["beta"]},
            {"category": "Three", "items": ["gamma"]},
        ],
    )
    result_more = run_patcher(src, replacements_more, out_more)
    if result_more.returncode != 0:
        raise AssertionError(f"multi-paragraph add case failed:\n{result_more.stdout}\n{result_more.stderr}")

    doc_more = Document(str(out_more))
    skills_more = skill_paragraphs(doc_more)
    if len(skills_more) != 3:
        raise AssertionError(f"expected cloned third Skills paragraph, got {len(skills_more)}")
    if [paragraph.text for paragraph in skills_more] != ["One: alpha", "Two: beta", "Three: gamma"]:
        raise AssertionError(f"unexpected multi-paragraph Skills text: {[p.text for p in skills_more]!r}")
    if len({paragraph.style.name for paragraph in skills_more}) != 1:
        raise AssertionError("cloned Skills paragraphs did not preserve source paragraph style")

    replacements_fewer = tmp_path / "multi_fewer_replacements.json"
    out_fewer = tmp_path / "multi_fewer_out.docx"
    write_replacements(replacements_fewer, [{"category": "Only", "items": ["delta"]}])
    result_fewer = run_patcher(src, replacements_fewer, out_fewer)
    if result_fewer.returncode != 0:
        raise AssertionError(f"multi-paragraph remove case failed:\n{result_fewer.stdout}\n{result_fewer.stderr}")

    doc_fewer = Document(str(out_fewer))
    skills_fewer = skill_paragraphs(doc_fewer)
    if len(skills_fewer) != 1 or skills_fewer[0].text != "Only: delta":
        raise AssertionError(f"extra old Skills paragraphs were not removed: {[p.text for p in skills_fewer]!r}")


def assert_malformed_structured_skills_fail(tmp_path: Path) -> None:
    replacements = tmp_path / "bad_replacements.json"
    out = tmp_path / "bad_out.docx"
    write_replacements(replacements, [{"category": "Broken", "items": ["valid", 42]}])
    result = run_patcher(REPO_ROOT / MASTER_NAME, replacements, out)
    if result.returncode == 0:
        raise AssertionError("malformed structured Skills unexpectedly succeeded")
    combined = f"{result.stdout}\n{result.stderr}"
    if "skills[0].items[1] must be a string" not in combined:
        raise AssertionError(f"malformed structured Skills failed for the wrong reason:\n{combined}")
    if out.exists():
        raise AssertionError("malformed structured Skills generated an output DOCX")


def assert_preformatted_skills_string_still_works(tmp_path: Path) -> None:
    replacements = tmp_path / "string_replacements.json"
    out = tmp_path / "string_out.docx"
    write_replacements(
        replacements,
        "String Category: one, two\nSecond String Category: three, four",
    )
    result = run_patcher(REPO_ROOT / MASTER_NAME, replacements, out)
    if result.returncode != 0:
        raise AssertionError(f"preformatted Skills string failed:\n{result.stdout}\n{result.stderr}")
    output_skills = skill_paragraphs(Document(str(out)))
    if len(output_skills) != 1:
        raise AssertionError(f"expected preformatted Skills string to render in one source paragraph, got {len(output_skills)}")
    if "String Category: one, two" not in output_skills[0].text:
        raise AssertionError("preformatted Skills string first line missing")
    if "Second String Category: three, four" not in output_skills[0].text:
        raise AssertionError("preformatted Skills string second line missing")


def main() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        assert_single_paragraph_structured_skills(tmp_path)
        assert_multi_paragraph_add_and_remove(tmp_path)
        assert_malformed_structured_skills_fail(tmp_path)
        assert_preformatted_skills_string_still_works(tmp_path)

    print("OK: structured Skills replacement preserves master formatting and validates malformed input.")


if __name__ == "__main__":
    main()
