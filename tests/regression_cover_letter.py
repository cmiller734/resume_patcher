#!/usr/bin/env python3
"""Regression checks for deterministic cover letter rendering."""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
MASTER_RESUME = "Caleb Miller Master Resume.docx"
COVER_LETTER = "Caleb Miller Cover Letter.docx"
PACKAGE_ROOT = "resume_patcher"


def run_patcher(args: list[str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(REPO_ROOT / "resume_patcher.py"), *args],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
    )


def write_empty_replacements(path: Path) -> None:
    path.write_text("{}", encoding="utf-8")


def write_cover_json(path: Path, paragraphs: list) -> None:
    path.write_text(
        json.dumps({"cover_letter": {"paragraphs": paragraphs}}, indent=2),
        encoding="utf-8",
    )


def make_legacy_cover_source(path: Path) -> None:
    shutil.copy2(REPO_ROOT / COVER_LETTER, path)
    doc = Document(str(path))
    for idx, paragraph in enumerate(doc.paragraphs):
        paragraph.text = f"Legacy Company stale cover letter paragraph {idx}"
    doc.save(str(path))


def non_empty_paragraph_texts(path: Path) -> list[str]:
    doc = Document(str(path))
    return [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]


def resume_header_texts(path: Path = REPO_ROOT / MASTER_RESUME) -> list[str]:
    doc = Document(str(path))
    texts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() and paragraph.style is not None and paragraph.style.name.lower() == "heading 1":
            break
        if paragraph.text.strip():
            texts.append(paragraph.text)
    return texts


def assert_valid_direct_cover_letter(tmp_path: Path) -> None:
    cover_src = tmp_path / COVER_LETTER
    cover_out = tmp_path / "tailored_cover.docx"
    resume_out = tmp_path / "tailored_resume.docx"
    replacements = tmp_path / "replacements.json"
    cover_json = tmp_path / "cover_letter.json"
    make_legacy_cover_source(cover_src)
    write_empty_replacements(replacements)

    paragraphs = [
        "Dear Hiring Team,",
        "I am excited to apply because this role matches my application support background.",
        "My experience spans enterprise SaaS troubleshooting, customer advising, and implementation support.",
        "Sincerely\nCaleb Miller",
    ]
    write_cover_json(cover_json, paragraphs)

    result = run_patcher(
        [
            "--src",
            str(REPO_ROOT / MASTER_RESUME),
            "--replacements",
            str(replacements),
            "--out",
            str(resume_out),
            "--cover-letter-docx",
            str(cover_src),
            "--cover-letter-json",
            str(cover_json),
            "--cover-letter-out",
            str(cover_out),
        ]
    )
    if result.returncode != 0:
        raise AssertionError(f"cover letter direct render failed:\n{result.stdout}\n{result.stderr}")

    actual = non_empty_paragraph_texts(cover_out)
    expected_header = resume_header_texts()
    if actual[: len(expected_header)] != expected_header:
        raise AssertionError(f"cover letter header did not match resume header: {actual[:len(expected_header)]!r}")
    if actual[len(expected_header) :] != paragraphs:
        raise AssertionError(f"cover letter paragraphs were not rendered in order: {actual!r}")
    combined = "\n".join(actual)
    if "Legacy Company" in combined:
        raise AssertionError("old cover letter body text remained in the output")
    if not resume_out.exists() or resume_out.stat().st_size <= 0:
        raise AssertionError("resume output was not generated alongside cover letter output")


def assert_bad_cover_json_fails(tmp_path: Path) -> None:
    cover_src = tmp_path / COVER_LETTER
    cover_out = tmp_path / "bad_cover.docx"
    resume_out = tmp_path / "bad_resume.docx"
    replacements = tmp_path / "replacements.json"
    cover_json = tmp_path / "bad_cover_letter.json"
    make_legacy_cover_source(cover_src)
    write_empty_replacements(replacements)
    write_cover_json(cover_json, ["Dear Hiring Team,", 42])

    result = run_patcher(
        [
            "--src",
            str(REPO_ROOT / MASTER_RESUME),
            "--replacements",
            str(replacements),
            "--out",
            str(resume_out),
            "--cover-letter-docx",
            str(cover_src),
            "--cover-letter-json",
            str(cover_json),
            "--cover-letter-out",
            str(cover_out),
        ]
    )
    if result.returncode == 0:
        raise AssertionError("malformed cover letter JSON unexpectedly succeeded")
    combined = f"{result.stdout}\n{result.stderr}"
    if "cover_letter.paragraphs[1] must be a string" not in combined:
        raise AssertionError(f"malformed cover letter failed for the wrong reason:\n{combined}")
    if cover_out.exists() or resume_out.exists():
        raise AssertionError("malformed cover letter JSON generated output")


def assert_mechanical_validation_fails(tmp_path: Path) -> None:
    cases = [
        ("double_hyphen", ["Dear Hiring Team,", "This has -- a double hyphen."]),
        ("placeholder", ["Dear {{COMPANY}},", "This still has a placeholder."]),
    ]

    for label, paragraphs in cases:
        cover_src = tmp_path / f"{label}_{COVER_LETTER}"
        cover_out = tmp_path / f"{label}_cover.docx"
        resume_out = tmp_path / f"{label}_resume.docx"
        replacements = tmp_path / f"{label}_replacements.json"
        cover_json = tmp_path / f"{label}_cover_letter.json"
        make_legacy_cover_source(cover_src)
        write_empty_replacements(replacements)
        write_cover_json(cover_json, paragraphs)

        result = run_patcher(
            [
                "--src",
                str(REPO_ROOT / MASTER_RESUME),
                "--replacements",
                str(replacements),
                "--out",
                str(resume_out),
                "--cover-letter-docx",
                str(cover_src),
                "--cover-letter-json",
                str(cover_json),
                "--cover-letter-out",
                str(cover_out),
            ]
        )
        if result.returncode == 0:
            raise AssertionError(f"{label} cover letter validation unexpectedly succeeded")
        if cover_out.exists() or resume_out.exists():
            raise AssertionError(f"{label} validation generated output")


def write_package_replacements(path: Path) -> None:
    path.write_text("{}", encoding="utf-8")


def build_cover_package(package_dir: Path, zip_path: Path, cover_out: Path) -> None:
    project_dir = package_dir / PACKAGE_ROOT
    project_dir.mkdir(parents=True, exist_ok=True)
    for name in [
        MASTER_RESUME,
        COVER_LETTER,
        "resume_patcher.py",
        "resume_preferences.json",
        "cover_letter_preferences.json",
        "requirements.txt",
        "README.md",
        "CHATBOT_POLICY.md",
        "COVER_LETTER_POLICY.md",
    ]:
        shutil.copy2(REPO_ROOT / name, project_dir / name)

    write_package_replacements(project_dir / "replacements.json")
    write_cover_json(
        project_dir / "cover_letter.json",
        [
            "Dear Hiring Team,",
            "Package mode writes this cover letter outside temp extraction.",
            "Sincerely\nCaleb Miller",
        ],
    )
    manifest = json.loads((REPO_ROOT / "manifest.json").read_text(encoding="utf-8"))
    manifest["output_path"] = str(package_dir / "tailored_resume.docx")
    manifest["cover_letter"]["output_path"] = str(cover_out)
    (project_dir / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    with ZipFile(zip_path, "w") as zf:
        for path in package_dir.rglob("*"):
            zf.write(path, path.relative_to(package_dir))


def assert_package_cover_output_persists(tmp_path: Path) -> None:
    package_dir = tmp_path / "package"
    zip_path = tmp_path / "cover_package.zip"
    cover_out = tmp_path / "package_cover.docx"
    build_cover_package(package_dir, zip_path, cover_out)

    result = run_patcher(["--package", str(zip_path)])
    if result.returncode != 0:
        raise AssertionError(f"package cover letter render failed:\n{result.stdout}\n{result.stderr}")
    actual = non_empty_paragraph_texts(cover_out)
    expected_header = resume_header_texts()
    expected_body = [
        "Dear Hiring Team,",
        "Package mode writes this cover letter outside temp extraction.",
        "Sincerely\nCaleb Miller",
    ]
    if actual[: len(expected_header)] != expected_header:
        raise AssertionError(f"package cover header did not match resume header: {actual[:len(expected_header)]!r}")
    if actual[len(expected_header) :] != expected_body:
        raise AssertionError(f"package cover output text mismatch: {actual!r}")


def main() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        assert_valid_direct_cover_letter(tmp_path)
        assert_bad_cover_json_fails(tmp_path)
        assert_mechanical_validation_fails(tmp_path)
        assert_package_cover_output_persists(tmp_path)

    print("OK: cover letter JSON rendering clears old body text and validates deterministic output.")


if __name__ == "__main__":
    main()
