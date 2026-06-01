#!/usr/bin/env python3
"""
resume_patcher.py

Deterministic DOCX resume patcher:
- preserve source DOCX formatting/layout as the source of truth
- patch text from external replacements JSON
- avoid rebuilding DOCX from scratch
"""
from __future__ import annotations

import argparse
import json
import re
import tempfile
from zipfile import BadZipFile, ZipFile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


BULLET_PREFIX_RE = re.compile(r"^\s*(?:\u2022|-|\*)\s+")
DATE_TOKEN_RE = re.compile(
    r"\b(?:jan\.?|january|feb\.?|february|mar\.?|march|apr\.?|april|may|jun\.?|june|"
    r"jul\.?|july|aug\.?|august|sep\.?|sept\.?|september|oct\.?|october|nov\.?|"
    r"november|dec\.?|december|present|\d{4})\b",
    re.IGNORECASE,
)
LOCATION_TOKEN_RE = re.compile(
    r"\b(?:remote|hybrid|onsite|on-site|[A-Z][a-z]+,\s*[A-Z]{2}|[A-Z]{2}\s*/\s*remote)\b",
    re.IGNORECASE,
)

SUPPORTED_STYLE_SOURCES = {
    "section_heading",
    "role_heading",
    "company_heading",
    "date_location",
    "normal",
    "bullet",
    "keep",
}
NON_SECTION_STYLE_SOURCES = {"normal", "role_heading", "company_heading", "date_location"}
CANONICAL_MASTER_RESUME_NAME = "Caleb Miller Master Resume.docx"
CANONICAL_COVER_LETTER_NAME = "Caleb Miller Cover Letter.docx"
FORBIDDEN_BASE_RESUME_FILENAMES = {
    "Caleb Miller - New Resume.docx",
    "Caleb Miller - Application Systems Engineer Resume.docx",
    "Caleb Miller - New Dev resume.docx",
}
DEFAULT_PACKAGE_MANIFEST = "manifest.json"
REQUIRED_PACKAGE_ROOT = "resume_patcher"
EXPECTED_OUTPUT_SECTION_HEADERS = ("SUMMARY", "SKILLS", "WORK EXPERIENCE", "EDUCATION")
ROLE_KEYWORDS = (
    "engineer",
    "developer",
    "architect",
    "consultant",
    "support",
    "services",
    "founder",
    "lead",
    "manager",
    "intern",
)
COVER_LETTER_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


@dataclass(frozen=True)
class PatchPara:
    text: str
    style_source: str = "normal"


@dataclass(frozen=True)
class SkillLine:
    text: str


@dataclass(frozen=True)
class CoverLetterRun:
    docx_path: Path
    json_path: Path
    output_path: Path


def warn(message: str) -> None:
    print(f"WARNING: {message}")


def ensure_non_empty_file(path: Path, label: str) -> None:
    if not path.exists():
        raise FileNotFoundError(f"{label} not found: {path}")
    if not path.is_file():
        raise FileNotFoundError(f"{label} is not a file: {path}")
    if path.stat().st_size <= 0:
        raise ValueError(f"{label} is empty: {path}")


def resolve_inside_root(root: Path, relative_path: str, label: str) -> Path:
    if not isinstance(relative_path, str) or not relative_path.strip():
        raise ValueError(f"{label} must be a non-empty relative path")

    candidate = Path(relative_path)
    if candidate.is_absolute():
        raise ValueError(f"{label} must be relative to the package root, got absolute path: {relative_path!r}")

    resolved_root = root.resolve()
    resolved = (resolved_root / candidate).resolve()
    if not resolved.is_relative_to(resolved_root):
        raise ValueError(f"{label} escapes the package root: {relative_path!r}")
    return resolved


def validate_base_resume_path(path: Path) -> None:
    if path.name in FORBIDDEN_BASE_RESUME_FILENAMES:
        raise ValueError(
            f"Forbidden base resume selected: {path.name!r}. "
            f"The only valid base resume is {CANONICAL_MASTER_RESUME_NAME!r}."
        )
    if path.name != CANONICAL_MASTER_RESUME_NAME:
        raise ValueError(
            f"Invalid base resume selected: {path.name!r}. "
            f"The only valid base resume is {CANONICAL_MASTER_RESUME_NAME!r}; no fallback resumes are allowed."
        )
    ensure_non_empty_file(path, "Master resume")


def normalized_doc_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip()).upper()


def validate_output_docx(path: Path) -> None:
    ensure_non_empty_file(path, "Output DOCX")
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Output DOCX could not be opened by python-docx: {path}") from exc

    paragraph_texts = [normalized_doc_text(p.text) for p in doc.paragraphs if p.text.strip()]
    missing = [
        header
        for header in EXPECTED_OUTPUT_SECTION_HEADERS
        if not any(normalized_doc_text(header) == text for text in paragraph_texts)
    ]
    if missing:
        raise ValueError(f"Output DOCX is missing expected resume section header(s): {', '.join(missing)}")


def is_ignored_package_artifact(path: str) -> bool:
    parts = [part for part in Path(path).parts if part not in {"", "."}]
    if not parts:
        return True
    for part in parts:
        if part == "__MACOSX" or part == ".DS_Store" or part.startswith("._") or part.startswith("~$"):
            return True
    return False


def safe_extract_zip(zip_path: Path, destination: Path) -> None:
    ensure_non_empty_file(zip_path, "Package ZIP")
    try:
        with ZipFile(zip_path) as zf:
            bad_member = zf.testzip()
            if bad_member is not None:
                raise ValueError(f"ZIP validation failed at member: {bad_member}")
            members = [info for info in zf.infolist() if info.filename and not info.is_dir()]
            if not members:
                raise ValueError(f"Package ZIP contains no files: {zip_path}")

            root = destination.resolve()
            for info in zf.infolist():
                target = (root / info.filename).resolve()
                if not target.is_relative_to(root):
                    raise ValueError(f"Unsafe ZIP member path escapes package root: {info.filename!r}")
            zf.extractall(destination)
    except BadZipFile as exc:
        raise ValueError(f"Package ZIP could not be opened: {zip_path}") from exc


def resolve_package_root(extraction_root: Path) -> Path:
    real_top_level = {
        Path(path).parts[0]
        for path in (str(p.relative_to(extraction_root)) for p in extraction_root.rglob("*"))
        if not is_ignored_package_artifact(path)
    }
    if REQUIRED_PACKAGE_ROOT not in real_top_level:
        raise FileNotFoundError(f"Required package root folder missing: {REQUIRED_PACKAGE_ROOT}/")

    unexpected = sorted(name for name in real_top_level if name != REQUIRED_PACKAGE_ROOT)
    if unexpected:
        raise ValueError(
            f"Unexpected real top-level package item(s): {', '.join(unexpected)}. "
            f"The only real top-level project folder must be {REQUIRED_PACKAGE_ROOT}/."
        )

    package_root = extraction_root / REQUIRED_PACKAGE_ROOT
    if not package_root.exists() or not package_root.is_dir():
        raise FileNotFoundError(f"Required package root is not a directory: {REQUIRED_PACKAGE_ROOT}/")
    return package_root


def load_package_manifest(path: Path) -> dict[str, Any]:
    ensure_non_empty_file(path, "Package manifest")
    try:
        with path.open("r", encoding="utf-8") as f:
            manifest = json.load(f)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in package manifest: {exc}") from exc

    if not isinstance(manifest, dict):
        raise ValueError("Package manifest must be a top-level object")
    return manifest


def validate_package_required_files(root: Path, manifest: dict[str, Any]) -> None:
    required_files = manifest.get("required_files")
    if not isinstance(required_files, list) or not required_files:
        raise ValueError("Package manifest must include a non-empty 'required_files' list")

    for idx, raw_path in enumerate(required_files):
        if not isinstance(raw_path, str):
            raise ValueError(f"Package manifest required_files[{idx}] must be a string")
        required_path = resolve_inside_root(root, raw_path, f"required_files[{idx}]")
        ensure_non_empty_file(required_path, f"Required package file {raw_path!r}")


def resolve_package_run_paths(
    root: Path,
    manifest: dict[str, Any],
    *,
    replacements_override: str | None,
    out_override: str | None,
) -> tuple[Path, Path, Path]:
    master_path_raw = manifest.get("master_resume_path") or manifest.get("master_resume")
    if not isinstance(master_path_raw, str):
        raise ValueError("Package manifest must include string 'master_resume_path'")

    src = resolve_inside_root(root, master_path_raw, "master_resume_path")
    validate_base_resume_path(src)

    replacements_raw = replacements_override or manifest.get("replacements_path") or manifest.get("replacements")
    if not isinstance(replacements_raw, str) or not replacements_raw.strip():
        raise ValueError("A replacements JSON path must be provided by --replacements or package manifest 'replacements_path'")

    replacements_candidate = Path(replacements_raw)
    if replacements_candidate.is_absolute():
        replacements = replacements_candidate
    else:
        replacements = resolve_inside_root(root, replacements_raw, "replacements_path")
    ensure_non_empty_file(replacements, "Replacements JSON")

    out_raw = out_override or manifest.get("output_path") or "tailored_resume.docx"
    if not isinstance(out_raw, str) or not out_raw.strip():
        raise ValueError("Output path must be a non-empty string")
    out = Path(out_raw)
    if not out.is_absolute():
        out = (Path.cwd() / out).resolve()

    return src, replacements, out


def resolve_package_file_override(root: Path, raw_path: str, label: str) -> Path:
    candidate = Path(raw_path)
    if candidate.is_absolute():
        return candidate
    return resolve_inside_root(root, raw_path, label)


def resolve_output_path(raw_path: str, label: str) -> Path:
    if not isinstance(raw_path, str) or not raw_path.strip():
        raise ValueError(f"{label} must be a non-empty string")

    output = Path(raw_path)
    if not output.is_absolute():
        output = (Path.cwd() / output).resolve()
    return output


def resolve_package_cover_letter_paths(
    root: Path,
    manifest: dict[str, Any],
    *,
    cover_letter_json_override: str | None,
    cover_letter_out_override: str | None,
    include_manifest_cover_letter: bool = True,
) -> CoverLetterRun | None:
    cover_manifest = manifest.get("cover_letter")
    if not include_manifest_cover_letter and cover_letter_json_override is None and cover_letter_out_override is None:
        return None
    if cover_manifest is None and cover_letter_json_override is None and cover_letter_out_override is None:
        return None
    if cover_manifest is None:
        raise ValueError("Package manifest must include 'cover_letter' when cover letter rendering is requested")
    if not isinstance(cover_manifest, dict):
        raise ValueError("Package manifest 'cover_letter' must be an object")

    docx_raw = cover_manifest.get("docx_path")
    if not isinstance(docx_raw, str) or not docx_raw.strip():
        raise ValueError("Package manifest cover_letter.docx_path must be a non-empty string")
    docx_path = resolve_inside_root(root, docx_raw, "cover_letter.docx_path")
    ensure_non_empty_file(docx_path, "Cover letter DOCX")

    json_raw = cover_letter_json_override or cover_manifest.get("json_path")
    if not isinstance(json_raw, str) or not json_raw.strip():
        raise ValueError("Package manifest cover_letter.json_path must be a non-empty string")
    json_path = resolve_package_file_override(root, json_raw, "cover_letter.json_path")
    ensure_non_empty_file(json_path, "Cover letter JSON")

    output_raw = cover_letter_out_override or cover_manifest.get("output_path") or docx_raw
    output_path = resolve_output_path(output_raw, "cover_letter.output_path")

    for key in ("policy_path", "preferences_path"):
        support_raw = cover_manifest.get(key)
        if support_raw is not None:
            if not isinstance(support_raw, str) or not support_raw.strip():
                raise ValueError(f"Package manifest cover_letter.{key} must be a non-empty string when provided")
            support_path = resolve_inside_root(root, support_raw, f"cover_letter.{key}")
            ensure_non_empty_file(support_path, f"Cover letter support file {support_raw!r}")

    return CoverLetterRun(docx_path=docx_path, json_path=json_path, output_path=output_path)


def prepare_package_run_with_cover(
    package_path: Path,
    *,
    manifest_name: str,
    replacements_override: str | None,
    out_override: str | None,
    cover_letter_json_override: str | None,
    cover_letter_out_override: str | None,
    include_manifest_cover_letter: bool = True,
) -> tuple[tempfile.TemporaryDirectory, Path, Path, Path, CoverLetterRun | None]:
    temp_dir = tempfile.TemporaryDirectory(prefix="resume_patcher_package_")
    extraction_root = Path(temp_dir.name)
    try:
        safe_extract_zip(package_path, extraction_root)
        package_root = resolve_package_root(extraction_root)
        manifest_path = resolve_inside_root(package_root, manifest_name, "manifest")
        manifest = load_package_manifest(manifest_path)
        validate_package_required_files(package_root, manifest)
        src, replacements, out = resolve_package_run_paths(
            package_root,
            manifest,
            replacements_override=replacements_override,
            out_override=out_override,
        )
        cover_letter_run = resolve_package_cover_letter_paths(
            package_root,
            manifest,
            cover_letter_json_override=cover_letter_json_override,
            cover_letter_out_override=cover_letter_out_override,
            include_manifest_cover_letter=include_manifest_cover_letter,
        )
    except Exception:
        temp_dir.cleanup()
        raise
    return temp_dir, src, replacements, out, cover_letter_run


def prepare_package_run(
    package_path: Path,
    *,
    manifest_name: str,
    replacements_override: str | None,
    out_override: str | None,
) -> tuple[tempfile.TemporaryDirectory, Path, Path, Path]:
    temp_dir, src, replacements, out, _cover_letter_run = prepare_package_run_with_cover(
        package_path,
        manifest_name=manifest_name,
        replacements_override=replacements_override,
        out_override=out_override,
        cover_letter_json_override=None,
        cover_letter_out_override=None,
        include_manifest_cover_letter=False,
    )
    return temp_dir, src, replacements, out


def apply_resume_bullet_indent(paragraph: Paragraph) -> None:
    """Apply bullet indentation to a paragraph.

    This helper is retained for compatibility, but the patcher primarily preserves
    bullet formatting by cloning existing source DOCX bullet paragraphs.
    """
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)

    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)

    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")


def _copy_run_format(src_run, dst_run) -> None:
    """Copy formatting from src_run to dst_run without copying text."""
    if src_run is None:
        return
    if src_run._r.rPr is not None:
        dst_run._r.insert(0, deepcopy(src_run._r.rPr))


def _template_runs_by_line(template: Paragraph | None) -> list:
    """Return the first visible run on each visual line of a template paragraph."""
    if template is None:
        return []

    line_runs = [None]
    line_idx = 0

    for run in template.runs:
        for child in run._r.iterchildren():
            if child.tag == qn("w:t"):
                value = child.text or ""
                if value.replace("\n", "").strip() and line_runs[line_idx] is None:
                    line_runs[line_idx] = run
            elif child.tag in {qn("w:br"), qn("w:cr")}:
                line_idx += 1
                while len(line_runs) <= line_idx:
                    line_runs.append(None)

        extra_breaks = run.text.count("\n")
        has_xml_break = any(child.tag in {qn("w:br"), qn("w:cr")} for child in run._r.iterchildren())
        if extra_breaks and not has_xml_break:
            for _ in range(extra_breaks):
                line_idx += 1
                while len(line_runs) <= line_idx:
                    line_runs.append(None)

    fallback = next((r for r in line_runs if r is not None), None)
    filled = []
    previous = fallback
    for r in line_runs:
        if r is not None:
            previous = r
            filled.append(r)
        else:
            filled.append(previous)
    return filled


def set_paragraph_text_keep_format(paragraph: Paragraph, text: str, template: Paragraph | None = None) -> None:
    """Replace paragraph text while preserving paragraph style and run formatting."""
    template = template or paragraph
    line_format_runs = _template_runs_by_line(template)
    fallback_run = (
        template.runs[0]
        if template is not None and template.runs
        else (paragraph.runs[0] if paragraph.runs else None)
    )

    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    parts = text.split("\n")
    previous_run = None
    for i, part in enumerate(parts):
        if i > 0 and previous_run is not None:
            previous_run.add_break()

        run = paragraph.add_run(part)
        src_run = line_format_runs[i] if i < len(line_format_runs) and line_format_runs[i] is not None else fallback_run
        _copy_run_format(src_run, run)
        previous_run = run


def _template_line_run_segments(template: Paragraph) -> list[list[tuple[str, Any]]]:
    lines: list[list[tuple[str, Any]]] = [[]]
    for run in template.runs:
        parts = run.text.split("\n")
        for idx, part in enumerate(parts):
            if idx > 0:
                lines.append([])
            if part:
                lines[-1].append((part, run))
    return lines


def _skill_line_format_runs(template: Paragraph, line_idx: int) -> tuple[Any, Any]:
    lines = _template_line_run_segments(template)
    if not lines:
        return None, None

    segments = lines[min(line_idx, len(lines) - 1)]
    if not segments:
        fallback = template.runs[0] if template.runs else None
        return fallback, fallback

    combined_text = "".join(text for text, _run in segments)
    colon_idx = combined_text.find(":")
    label_run = next((run for text, run in segments if text.strip()), segments[0][1])
    rest_run = None

    if colon_idx >= 0:
        pos = 0
        for text, run in segments:
            next_pos = pos + len(text)
            if pos <= colon_idx < next_pos:
                label_run = run
            if next_pos > colon_idx + 1 and text[max(0, colon_idx + 1 - pos) :].strip():
                rest_run = run
                break
            pos = next_pos

    if rest_run is None:
        rest_run = next((run for text, run in segments if text.strip() and run is not label_run), label_run)

    return label_run, rest_run


def set_skills_text_keep_label_format(paragraph: Paragraph, text: str, template: Paragraph) -> None:
    """Replace skills text while preserving the template's label/rest run pattern."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    previous_run = None
    for line_idx, line in enumerate(text.split("\n")):
        if line_idx > 0:
            if previous_run is None:
                previous_run = paragraph.add_run("")
            previous_run.add_break()

        label_source, rest_source = _skill_line_format_runs(template, line_idx)
        colon_idx = line.find(":")
        if colon_idx > 0:
            label_run = paragraph.add_run(line[: colon_idx + 1])
            _copy_run_format(label_source, label_run)
            previous_run = label_run

            rest = line[colon_idx + 1 :]
            if rest:
                rest_run = paragraph.add_run(rest)
                _copy_run_format(rest_source, rest_run)
                previous_run = rest_run
        else:
            run = paragraph.add_run(line)
            _copy_run_format(rest_source or label_source, run)
            previous_run = run


def insert_skill_paragraph_after(paragraph: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_p = deepcopy(template._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = template.style
    set_skills_text_keep_label_format(inserted, text, template)
    return inserted


def insert_paragraph_after(paragraph: Paragraph, text: str, template: Paragraph) -> Paragraph:
    """Insert a paragraph after another paragraph, copying style and paragraph properties."""
    new_p = deepcopy(template._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = template.style
    set_paragraph_text_keep_format(inserted, text, template)
    return inserted


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def copy_paragraph_properties(destination: Paragraph, source: Paragraph) -> None:
    existing_ppr = destination._p.pPr
    if existing_ppr is not None:
        destination._p.remove(existing_ppr)
    if source._p.pPr is not None:
        destination._p.insert(0, deepcopy(source._p.pPr))
    destination.style = source.style


def load_cover_letter_paragraphs(path: Path) -> list[str]:
    ensure_non_empty_file(path, "Cover letter JSON")
    try:
        with path.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in cover letter file: {exc}") from exc

    if not isinstance(data, dict):
        raise ValueError("Cover letter JSON must be a top-level object")
    cover_letter = data.get("cover_letter")
    if not isinstance(cover_letter, dict):
        raise ValueError("Cover letter JSON must include object 'cover_letter'")

    paragraphs = cover_letter.get("paragraphs")
    if not isinstance(paragraphs, list) or not paragraphs:
        raise ValueError("cover_letter.paragraphs must be a non-empty list of strings")

    validated: list[str] = []
    for idx, paragraph in enumerate(paragraphs):
        if not isinstance(paragraph, str):
            raise ValueError(f"cover_letter.paragraphs[{idx}] must be a string")
        if not paragraph.strip():
            raise ValueError(f"cover_letter.paragraphs[{idx}] must not be empty")
        if "--" in paragraph:
            raise ValueError(f"cover_letter.paragraphs[{idx}] contains double hyphen '--'")
        placeholder = COVER_LETTER_PLACEHOLDER_RE.search(paragraph)
        if placeholder is not None:
            raise ValueError(
                f"cover_letter.paragraphs[{idx}] contains unreplaced placeholder {placeholder.group(0)!r}"
            )
        validated.append(paragraph)

    return validated


def validate_cover_letter_input(run: CoverLetterRun) -> list[str]:
    ensure_non_empty_file(run.docx_path, "Cover letter DOCX")
    try:
        doc = Document(str(run.docx_path))
    except Exception as exc:
        raise ValueError(f"Cover letter DOCX could not be opened by python-docx: {run.docx_path}") from exc
    if not doc.paragraphs:
        raise ValueError("Cover letter DOCX contains no body paragraphs to use as a formatting template")

    return load_cover_letter_paragraphs(run.json_path)


def validate_cover_letter_output_docx(path: Path, expected_paragraphs: Sequence[str]) -> None:
    ensure_non_empty_file(path, "Cover letter output DOCX")
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Cover letter output DOCX could not be opened by python-docx: {path}") from exc

    actual_paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    if actual_paragraphs != list(expected_paragraphs):
        raise ValueError("Cover letter output paragraphs did not match cover_letter.json in order")
    for idx, paragraph in enumerate(actual_paragraphs):
        if "--" in paragraph:
            raise ValueError(f"Cover letter output paragraph {idx} contains double hyphen '--'")
        placeholder = COVER_LETTER_PLACEHOLDER_RE.search(paragraph)
        if placeholder is not None:
            raise ValueError(f"Cover letter output paragraph {idx} contains unreplaced placeholder {placeholder.group(0)!r}")


def render_cover_letter_docx(run: CoverLetterRun, paragraphs: Sequence[str]) -> None:
    try:
        doc = Document(str(run.docx_path))
    except Exception as exc:
        raise ValueError(f"Cover letter DOCX could not be opened by python-docx: {run.docx_path}") from exc
    if not doc.paragraphs:
        raise ValueError("Cover letter DOCX contains no body paragraphs to use as a formatting template")

    template_source = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip()), doc.paragraphs[0])
    template = Paragraph(deepcopy(template_source._p), template_source._parent)
    first = doc.paragraphs[0]

    for paragraph in list(doc.paragraphs[1:]):
        delete_paragraph(paragraph)

    copy_paragraph_properties(first, template)
    set_paragraph_text_keep_format(first, paragraphs[0], template)
    cursor = first
    for text in paragraphs[1:]:
        cursor = insert_paragraph_after(cursor, text, template)

    run.output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(run.output_path))
    validate_cover_letter_output_docx(run.output_path, paragraphs)


def find_para_index(doc: Document, exact_text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == exact_text.strip():
            return i
    raise ValueError(f"Could not find paragraph with exact text: {exact_text!r}")


def find_para_contains(doc: Document, needle: str, start: int = 0, *, case_sensitive: bool = False) -> int:
    target = needle if case_sensitive else needle.lower()
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        haystack = p.text if case_sensitive else p.text.lower()
        if target in haystack:
            return i
    raise ValueError(f"Could not find paragraph containing: {needle!r}")


def replace_block(
    doc: Document,
    start_idx: int,
    end_idx_exclusive: int,
    items: Sequence[PatchPara],
    *,
    style_templates: dict[str, Paragraph],
) -> None:
    """Replace a paragraph block [start_idx, end_idx_exclusive) with items."""
    if not items:
        raise ValueError("replace_block requires at least one item")

    first = doc.paragraphs[start_idx]
    end_anchor = doc.paragraphs[end_idx_exclusive]._p

    template_xml_by_style = {
        style_source: deepcopy(template._p)
        for style_source, template in style_templates.items()
    }
    template_style_name_by_style = {
        style_source: paragraph_style_name(template)
        for style_source, template in style_templates.items()
    }

    def make_template(style_source: str, current: Paragraph | None = None) -> Paragraph:
        if style_source == "keep":
            return current or Paragraph(deepcopy(template_xml_by_style["normal"]), first._parent)

        def safe_normal_template() -> Paragraph:
            fallback = Paragraph(deepcopy(template_xml_by_style["normal"]), first._parent)
            if paragraph_is_heading_1(fallback):
                warn("safe 'normal' fallback resolved to Heading 1; forcing Word Normal style")
                fallback.style = "Normal"
            return fallback

        if style_source not in template_xml_by_style:
            warn(f"style {style_source!r} is not available; using 'normal'")
            style_source = "normal"

        template = Paragraph(deepcopy(template_xml_by_style[style_source]), first._parent)
        if style_source in NON_SECTION_STYLE_SOURCES and paragraph_is_heading_1(template):
            warn(
                f"refusing to apply Heading 1 template to {style_source!r}; "
                "using safe 'normal' template"
            )
            template = safe_normal_template()
        elif style_source in NON_SECTION_STYLE_SOURCES and "heading 1" in template_style_name_by_style.get(style_source, "").lower():
            warn(
                f"refusing to apply Heading 1 template to {style_source!r}; "
                "using safe 'normal' template"
            )
            template = safe_normal_template()
        return template

    first.style = make_template(items[0].style_source, first).style
    set_paragraph_text_keep_format(first, items[0].text, make_template(items[0].style_source, first))

    while first._p.getnext() is not None and first._p.getnext() is not end_anchor:
        delete_paragraph(Paragraph(first._p.getnext(), first._parent))

    cursor = first
    for item in items[1:]:
        cursor = insert_paragraph_after(cursor, item.text, make_template(item.style_source))


def paragraph_is_bullet(paragraph: Paragraph) -> bool:
    pPr = paragraph._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
    if "bullet" in style_name:
        return True
    return bool(BULLET_PREFIX_RE.match(paragraph.text))


def paragraph_style_name(paragraph: Paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def paragraph_style_id(paragraph: Paragraph) -> str:
    return paragraph.style.style_id if paragraph.style is not None else ""


def paragraph_has_heading_level(paragraph: Paragraph, level: int) -> bool:
    style_name = paragraph_style_name(paragraph).lower().replace(" ", "")
    style_id = paragraph_style_id(paragraph).lower().replace(" ", "")
    return style_name == f"heading{level}" or style_id == f"heading{level}"


def paragraph_is_heading_1(paragraph: Paragraph) -> bool:
    return paragraph_has_heading_level(paragraph, 1)


def paragraph_is_any_heading(paragraph: Paragraph) -> bool:
    style_name = paragraph_style_name(paragraph).lower()
    style_id = paragraph_style_id(paragraph).lower()
    return "heading" in style_name or style_id.startswith("heading")


def paragraph_is_document_title_style(paragraph: Paragraph) -> bool:
    style_name = paragraph_style_name(paragraph).lower()
    style_id = paragraph_style_id(paragraph).lower()
    return style_name in {"title", "subtitle"} or style_id in {"title", "subtitle"}


def paragraph_looks_like_section_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False

    if paragraph_is_any_heading(paragraph) and paragraph_has_heading_level(paragraph, 1):
        return True

    has_letters = any(ch.isalpha() for ch in text)
    return has_letters and text.upper() == text and len(text) <= 80


def paragraph_looks_like_date_location(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or paragraph_is_bullet(paragraph) or paragraph_looks_like_section_heading(paragraph):
        return False
    if len(text) > 90:
        return False
    if any(word in text.lower() for word in ROLE_KEYWORDS):
        return False
    has_date = DATE_TOKEN_RE.search(text) is not None
    has_location = LOCATION_TOKEN_RE.search(text) is not None
    return has_date and (has_location or "|" in text or " - " in text or " – " in text)


def text_looks_like_date_location(text: str) -> bool:
    text = text.strip()
    if not text or len(text) > 120:
        return False
    has_date = DATE_TOKEN_RE.search(text) is not None
    has_location = LOCATION_TOKEN_RE.search(text) is not None
    return has_date and (has_location or "|" in text or " - " in text or " – " in text)


def paragraph_looks_like_company_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or paragraph_is_bullet(paragraph) or paragraph_looks_like_section_heading(paragraph):
        return False
    if paragraph_has_heading_level(paragraph, 2):
        return True
    if DATE_TOKEN_RE.search(text):
        return False
    return " | " in text or " – " in text or " - " in text


def paragraph_looks_like_role_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or paragraph_is_bullet(paragraph) or paragraph_looks_like_section_heading(paragraph):
        return False
    if paragraph_has_heading_level(paragraph, 3):
        return True
    lowered = text.lower()
    return any(word in lowered for word in ROLE_KEYWORDS) and DATE_TOKEN_RE.search(text) is not None


def paragraph_looks_like_combined_role_date_template(paragraph: Paragraph) -> bool:
    lines = [line.strip() for line in paragraph.text.splitlines() if line.strip()]
    return len(lines) >= 2 and "|" not in lines[0] and text_looks_like_date_location(lines[1])


def first_non_empty_paragraph(doc: Document) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip():
            return p
    return doc.paragraphs[0]


def first_matching_paragraph(doc: Document, predicate) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() and predicate(p):
            return p
    return None


def derive_line_template(template: Paragraph, line_idx: int) -> Paragraph:
    """Create a one-line template using a specific visual line from a source paragraph."""
    derived = Paragraph(deepcopy(template._p), template._parent)
    lines = template.text.splitlines() or [template.text]
    line_text = lines[line_idx] if line_idx < len(lines) else (lines[-1] if lines else "")
    line_runs = _template_runs_by_line(template)
    source_run = line_runs[line_idx] if line_idx < len(line_runs) else None
    if source_run is None and line_runs:
        source_run = line_runs[-1]

    for run in list(derived.runs):
        run._element.getparent().remove(run._element)

    run = derived.add_run(line_text)
    _copy_run_format(source_run, run)
    return derived


def first_safe_non_heading_paragraph(doc: Document) -> Paragraph:
    for p in doc.paragraphs:
        if (
            p.text.strip()
            and not paragraph_is_heading_1(p)
            and not paragraph_is_bullet(p)
            and not paragraph_is_document_title_style(p)
        ):
            return p
    return first_non_empty_paragraph(doc)


def paragraph_is_safe_normal_template(paragraph: Paragraph) -> bool:
    return (
        bool(paragraph.text.strip())
        and not paragraph_is_any_heading(paragraph)
        and not paragraph_is_document_title_style(paragraph)
        and not paragraph_is_bullet(paragraph)
        and not paragraph_looks_like_date_location(paragraph)
        and not paragraph_looks_like_company_heading(paragraph)
        and not paragraph_looks_like_role_heading(paragraph)
    )


def discover_style_templates(doc: Document) -> dict[str, Paragraph]:
    """Map semantic replacement styles to real paragraph templates in the source DOCX."""
    safe_non_heading = first_safe_non_heading_paragraph(doc)
    combined_role_date_template = first_matching_paragraph(doc, paragraph_looks_like_combined_role_date_template)

    templates: dict[str, Paragraph] = {}
    templates["section_heading"] = (
        first_matching_paragraph(doc, paragraph_looks_like_section_heading)
        or first_non_empty_paragraph(doc)
    )
    templates["normal"] = (
        first_matching_paragraph(doc, paragraph_is_safe_normal_template)
        or safe_non_heading
    )
    templates["bullet"] = (
        first_matching_paragraph(doc, paragraph_is_bullet)
        or templates["normal"]
    )
    templates["company_heading"] = (
        first_matching_paragraph(doc, lambda p: paragraph_has_heading_level(p, 2))
        or first_matching_paragraph(doc, paragraph_looks_like_company_heading)
        or templates["normal"]
    )
    templates["role_heading"] = (
        (derive_line_template(combined_role_date_template, 0) if combined_role_date_template is not None else None)
        or first_matching_paragraph(doc, lambda p: paragraph_has_heading_level(p, 3))
        or first_matching_paragraph(doc, paragraph_looks_like_role_heading)
        or templates["company_heading"]
        or templates["normal"]
    )
    templates["date_location"] = (
        (derive_line_template(combined_role_date_template, 1) if combined_role_date_template is not None else None)
        or first_matching_paragraph(doc, paragraph_looks_like_date_location)
        or templates["normal"]
    )

    for style_source in NON_SECTION_STYLE_SOURCES:
        template = templates[style_source]
        if paragraph_is_heading_1(template):
            warn(
                f"semantic style {style_source!r} resolved to Heading 1; "
                "using a non-section fallback instead"
            )
            templates[style_source] = safe_non_heading

    return templates


def infer_end_index(doc: Document, start_idx: int) -> int | None:
    start_para = doc.paragraphs[start_idx]
    start_style_id = start_para.style.style_id if start_para.style is not None else ""

    for i in range(start_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if not p.text.strip():
            continue

        if paragraph_looks_like_section_heading(p):
            return i

        style_id = p.style.style_id if p.style is not None else ""
        if style_id == start_style_id and not paragraph_is_bullet(p):
            return i

    return None


def parse_patch_paragraphs(raw_items: Any, context: str) -> list[PatchPara]:
    if not isinstance(raw_items, list):
        raise ValueError(f"{context} must be a list")

    parsed: list[PatchPara] = []
    for idx, raw in enumerate(raw_items):
        if isinstance(raw, str):
            style_source = "normal"
            text = raw
            if BULLET_PREFIX_RE.match(text):
                style_source = "bullet"
                text = BULLET_PREFIX_RE.sub("", text, count=1)
            parsed.append(PatchPara(text=text, style_source=style_source))
            continue

        if isinstance(raw, dict):
            text = raw.get("text")
            if not isinstance(text, str):
                raise ValueError(f"{context}[{idx}].text must be a string")

            style_source = raw.get("style", "normal")
            if style_source not in SUPPORTED_STYLE_SOURCES:
                warn(f"{context}[{idx}] has unknown style {style_source!r}; using 'normal'")
                style_source = "normal"

            if style_source == "bullet":
                text = BULLET_PREFIX_RE.sub("", text, count=1)

            parsed.append(PatchPara(text=text, style_source=style_source))
            continue

        raise ValueError(f"{context}[{idx}] must be either a string or an object")

    return parsed


def apply_paragraph_replacements(doc: Document, replacements: Any) -> int:
    if not replacements:
        return 0
    if not isinstance(replacements, list):
        raise ValueError("paragraph_replacements must be a list")

    applied = 0
    for idx, item in enumerate(replacements):
        if not isinstance(item, dict):
            warn(f"paragraph_replacements[{idx}] is not an object; skipping")
            continue

        find = item.get("find")
        replace = item.get("replace")
        if not isinstance(find, str) or not isinstance(replace, str):
            warn(f"paragraph_replacements[{idx}] requires string 'find' and 'replace'; skipping")
            continue

        try:
            para_idx = find_para_index(doc, find)
        except ValueError:
            warn(f"paragraph_replacements[{idx}] could not find exact text: {find!r}")
            continue

        target = doc.paragraphs[para_idx]
        set_paragraph_text_keep_format(target, replace, target)
        applied += 1

    return applied


def find_summary_paragraph_index(doc: Document, heading_text: str = "SUMMARY") -> int | None:
    try:
        heading_idx = find_para_contains(doc, heading_text)
    except ValueError:
        return None

    for i in range(heading_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            return i
    return None


def normalize_skill_items(raw_items: Any, context: str) -> str:
    if isinstance(raw_items, str):
        return raw_items.strip()

    if isinstance(raw_items, list):
        normalized: list[str] = []
        for idx, item in enumerate(raw_items):
            if not isinstance(item, str):
                raise ValueError(f"{context}[{idx}] must be a string")
            item = item.strip()
            if item:
                normalized.append(item)
        return ", ".join(normalized)

    raise ValueError(f"{context} must be a string or a list of strings")


def render_structured_skill_line(raw: Any, context: str) -> SkillLine:
    if isinstance(raw, str):
        text = raw.strip()
        if not text:
            raise ValueError(f"{context} must not be empty")
        return SkillLine(text=text)

    if not isinstance(raw, dict):
        raise ValueError(
            f"{context} must be a string or an object with string 'category' and string/list 'items'"
        )

    category = raw.get("category")
    if not isinstance(category, str) or not category.strip():
        raise ValueError(f"{context}.category must be a non-empty string")

    if "items" not in raw:
        raise ValueError(f"{context}.items is required and must be a string or a list of strings")

    items_text = normalize_skill_items(raw.get("items"), f"{context}.items")
    separator = " " if items_text else ""
    return SkillLine(text=f"{category.strip()}:{separator}{items_text}")


def normalize_skills_replacement(skills: Any) -> list[SkillLine]:
    if skills is None:
        return []

    if isinstance(skills, str):
        lines = [line.strip() for line in skills.splitlines() if line.strip()]
        if not lines:
            raise ValueError("skills string must not be empty")
        return [SkillLine(text=line) for line in lines]

    if isinstance(skills, dict):
        if "categories" in skills:
            categories = skills.get("categories")
            if not isinstance(categories, list) or not categories:
                raise ValueError("skills.categories must be a non-empty list")
            return [render_structured_skill_line(raw, f"skills.categories[{idx}]") for idx, raw in enumerate(categories)]

        if not skills:
            raise ValueError("skills object must not be empty")

        lines: list[SkillLine] = []
        for label, value in skills.items():
            if not isinstance(label, str) or not label.strip():
                raise ValueError("skills object keys must be non-empty category strings")
            items_text = normalize_skill_items(value, f"skills[{label!r}]")
            separator = " " if items_text else ""
            lines.append(SkillLine(text=f"{label.strip()}:{separator}{items_text}"))
        return lines

    if isinstance(skills, list):
        if not skills:
            raise ValueError("skills list must not be empty")
        return [render_structured_skill_line(raw, f"skills[{idx}]") for idx, raw in enumerate(skills)]

    raise ValueError(
        "skills must be an object mapping category names to strings/lists, "
        "or a list of strings/objects with 'category' and 'items'"
    )


def find_skills_section_bounds(doc: Document) -> tuple[int, int] | None:
    heading_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().upper() == "SKILLS":
            heading_idx = i
            break

    if heading_idx is None:
        return None

    end_idx = len(doc.paragraphs)
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        if paragraph.text.strip() and paragraph_looks_like_section_heading(paragraph):
            end_idx = i
            break

    return heading_idx, end_idx


def render_skills_lines(doc: Document, lines: list[SkillLine]) -> int:
    bounds = find_skills_section_bounds(doc)
    if bounds is None:
        raise ValueError("skills replacement provided but SKILLS section was not found")

    heading_idx, end_idx = bounds
    targets = [doc.paragraphs[i] for i in range(heading_idx + 1, end_idx) if doc.paragraphs[i].text.strip()]
    if not targets:
        raise ValueError("skills replacement provided but no existing Skills paragraph template was found")

    templates = [Paragraph(deepcopy(target._p), target._parent) for target in targets]
    rendered_lines = [line.text for line in lines]

    if len(targets) == 1:
        set_skills_text_keep_label_format(targets[0], "\n".join(rendered_lines), templates[0])
        return len(rendered_lines)

    cursor = targets[0]
    for idx, text in enumerate(rendered_lines):
        template = templates[min(idx, len(templates) - 1)]
        if idx < len(targets):
            target = targets[idx]
            set_skills_text_keep_label_format(target, text, template)
            cursor = target
        else:
            cursor = insert_skill_paragraph_after(cursor, text, template)

    for extra in targets[len(rendered_lines) :]:
        delete_paragraph(extra)

    return len(rendered_lines)


def apply_summary_replacement(doc: Document, summary: Any) -> int:
    if summary is None:
        return 0

    if isinstance(summary, str):
        target_idx = find_summary_paragraph_index(doc)
        if target_idx is None:
            warn("summary was provided but no SUMMARY heading/paragraph target was found")
            return 0

        target = doc.paragraphs[target_idx]
        set_paragraph_text_keep_format(target, summary, target)
        return 1

    if isinstance(summary, dict):
        text = summary.get("text") or summary.get("replace")
        if not isinstance(text, str):
            warn("summary object must include string 'text' (or 'replace'); skipping")
            return 0

        find_text = summary.get("find")
        if isinstance(find_text, str):
            try:
                para_idx = find_para_index(doc, find_text)
            except ValueError:
                warn(f"summary.find exact text not found: {find_text!r}")
                return 0
            target = doc.paragraphs[para_idx]
            set_paragraph_text_keep_format(target, text, target)
            return 1

        heading = summary.get("heading", "SUMMARY")
        if not isinstance(heading, str):
            warn("summary.heading must be a string when provided; using 'SUMMARY'")
            heading = "SUMMARY"
        target_idx = find_summary_paragraph_index(doc, heading)
        if target_idx is None:
            warn(f"summary heading not found: {heading!r}")
            return 0

        target = doc.paragraphs[target_idx]
        set_paragraph_text_keep_format(target, text, target)
        return 1

    warn("summary must be a string or object; skipping")
    return 0


def apply_skills_replacements(doc: Document, skills: Any) -> int:
    if skills is None:
        return 0

    lines = normalize_skills_replacement(skills)
    return render_skills_lines(doc, lines)


def apply_block_items(
    doc: Document,
    start_idx: int,
    end_idx: int,
    items: list[PatchPara],
    context: str,
    style_templates: dict[str, Paragraph],
) -> bool:
    if not items:
        warn(f"{context} had no replacement_paragraphs after parsing; skipping")
        return False

    if end_idx <= start_idx:
        warn(f"{context} has invalid range start={start_idx}, end={end_idx}; skipping")
        return False

    end_anchor_text = doc.paragraphs[end_idx].text.strip()
    if items and items[-1].text.strip() == end_anchor_text:
        items = items[:-1]
        if not items:
            warn(f"{context} only repeated the end heading anchor; skipping")
            return False

    replace_block(
        doc,
        start_idx,
        end_idx,
        items,
        style_templates=style_templates,
    )
    return True


def apply_named_blocks(
    doc: Document,
    raw_blocks: Any,
    *,
    label: str,
    style_templates: dict[str, Paragraph],
    default_use_heading_key: bool = False,
    require_end_heading: bool = False,
) -> int:
    if not raw_blocks:
        return 0
    if not isinstance(raw_blocks, list):
        raise ValueError(f"{label} must be a list")

    applied = 0
    for idx, block in enumerate(raw_blocks):
        context = f"{label}[{idx}]"
        if not isinstance(block, dict):
            warn(f"{context} is not an object; skipping")
            continue

        if default_use_heading_key:
            start_heading = block.get("start_heading") or block.get("heading")
        else:
            start_heading = block.get("start_heading")

        if not isinstance(start_heading, str) or not start_heading.strip():
            warn(f"{context} requires a non-empty 'start_heading' or 'heading'; skipping")
            continue

        replacement_paragraphs = block.get("replacement_paragraphs")
        if replacement_paragraphs is None:
            replacement_paragraphs = block.get("content")
        if replacement_paragraphs is None:
            warn(f"{context} is missing 'replacement_paragraphs' or 'content'; skipping")
            continue

        try:
            items = parse_patch_paragraphs(replacement_paragraphs, f"{context}.replacement_paragraphs")
        except ValueError as exc:
            warn(str(exc))
            continue

        try:
            start_idx = find_para_contains(doc, start_heading)
        except ValueError:
            warn(f"{context} start heading not found: {start_heading!r}")
            continue

        end_heading = block.get("end_heading")
        end_idx = None

        if isinstance(end_heading, str) and end_heading.strip():
            try:
                end_idx = find_para_contains(doc, end_heading, start=start_idx + 1)
            except ValueError:
                warn(f"{context} end heading not found: {end_heading!r}")
                continue
        elif require_end_heading:
            warn(f"{context} requires 'end_heading'; skipping")
            continue
        else:
            end_idx = infer_end_index(doc, start_idx)
            if end_idx is None:
                warn(f"{context} could not infer an end boundary; provide 'end_heading' to disambiguate")
                continue

        if apply_block_items(doc, start_idx, end_idx, items, context, style_templates):
            applied += 1

    return applied


def load_replacements(path: Path) -> dict[str, Any]:
    ensure_non_empty_file(path, "Replacements JSON")

    try:
        with path.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in replacements file: {exc}") from exc

    if not isinstance(data, dict):
        raise ValueError("Replacements JSON must be a top-level object")

    return data


def run_patch(src: Path, replacements_path: Path, out: Path) -> dict[str, int]:
    validate_base_resume_path(src)
    replacements = load_replacements(replacements_path)

    doc = Document(str(src))
    if not doc.paragraphs:
        raise ValueError("Source DOCX contains no paragraphs")
    style_templates = discover_style_templates(doc)

    applied_counts: dict[str, int] = {}

    applied_counts["paragraph_replacements"] = apply_paragraph_replacements(doc, replacements.get("paragraph_replacements"))
    applied_counts["summary"] = apply_summary_replacement(doc, replacements.get("summary"))
    applied_counts["skills"] = apply_skills_replacements(doc, replacements.get("skills"))
    applied_counts["experience_blocks"] = apply_named_blocks(
        doc,
        replacements.get("experience_blocks"),
        label="experience_blocks",
        style_templates=style_templates,
        default_use_heading_key=True,
        require_end_heading=False,
    )
    applied_counts["section_replacements"] = apply_named_blocks(
        doc,
        replacements.get("section_replacements"),
        label="section_replacements",
        style_templates=style_templates,
        default_use_heading_key=True,
        require_end_heading=False,
    )
    applied_counts["block_replacements"] = apply_named_blocks(
        doc,
        replacements.get("block_replacements"),
        label="block_replacements",
        style_templates=style_templates,
        default_use_heading_key=False,
        require_end_heading=True,
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    validate_output_docx(out)
    return applied_counts


def main() -> None:
    parser = argparse.ArgumentParser(description="Patch a resume DOCX using external replacements JSON while preserving formatting.")
    parser.add_argument("--package", help="ZIP package containing the master resume and package manifest")
    parser.add_argument("--manifest", default=DEFAULT_PACKAGE_MANIFEST, help="Package-relative manifest path for --package mode")
    parser.add_argument("--src", help="Source DOCX resume for direct mode")
    parser.add_argument("--replacements", help="Replacements JSON path")
    parser.add_argument("--out", help="Output DOCX path")
    parser.add_argument("--cover-letter-docx", help="Cover letter DOCX template/path for direct mode")
    parser.add_argument("--cover-letter-json", help="Cover letter JSON path")
    parser.add_argument("--cover-letter-out", help="Cover letter output DOCX path")
    args = parser.parse_args()

    package_temp = None
    cover_letter_run = None
    cover_letter_paragraphs = None
    try:
        if args.package:
            if args.src is not None:
                raise ValueError("--src is not allowed in package mode; the package manifest selects the base resume")
            if args.cover_letter_docx is not None:
                raise ValueError("--cover-letter-docx is not allowed in package mode; the package manifest selects the cover letter DOCX")
            package_temp, src, replacements_path, out, cover_letter_run = prepare_package_run_with_cover(
                Path(args.package),
                manifest_name=args.manifest,
                replacements_override=args.replacements,
                out_override=args.out,
                cover_letter_json_override=args.cover_letter_json,
                cover_letter_out_override=args.cover_letter_out,
            )
        else:
            missing = [name for name in ("src", "replacements", "out") if getattr(args, name) is None]
            if missing:
                raise ValueError(
                    "Direct mode requires --src, --replacements, and --out. "
                    "Package mode requires --package."
                )
            src = Path(args.src)
            replacements_path = Path(args.replacements)
            out = Path(args.out)
            if args.cover_letter_docx is not None or args.cover_letter_json is not None or args.cover_letter_out is not None:
                if args.cover_letter_json is None:
                    raise ValueError("Direct mode cover letter rendering requires --cover-letter-json")
                cover_letter_docx = Path(args.cover_letter_docx or CANONICAL_COVER_LETTER_NAME)
                cover_letter_out = Path(args.cover_letter_out) if args.cover_letter_out else cover_letter_docx
                cover_letter_run = CoverLetterRun(
                    docx_path=cover_letter_docx,
                    json_path=Path(args.cover_letter_json),
                    output_path=cover_letter_out,
                )

        if cover_letter_run is not None:
            cover_letter_paragraphs = validate_cover_letter_input(cover_letter_run)
        applied_counts = run_patch(src, replacements_path, out)
        if cover_letter_run is not None:
            render_cover_letter_docx(cover_letter_run, cover_letter_paragraphs)
    finally:
        if package_temp is not None:
            package_temp.cleanup()

    applied_summary = ", ".join(f"{k}={v}" for k, v in applied_counts.items())
    print(f"Wrote {out}")
    if cover_letter_run is not None:
        print(f"Wrote cover letter {cover_letter_run.output_path}")
    print(f"Applied operations: {applied_summary}")


if __name__ == "__main__":
    main()
